#!/usr/bin/env python3
"""Merge docx files into one using docxcompose.

Usage: python3 merge_docx.py <host.docx> <annex1.docx> [<annex2.docx> ...] <output.docx>
First arg is the host, LAST arg is the output, everything between is appended
in order. The original 3-arg form (host annex output) still works unchanged.

Every appended annex is preceded by an explicit PAGE BREAK so merged pieces
never run together on the same page (David, Aug 1: the Visual report's cover
page flowed straight into the first section because docxcompose appends
content with no break; "Pages CANNOT run together"). Each piece - cover,
building, location, section chunk, and the Visual annex of a Final report -
starts at the top of its own page.

Exit code 0 on success; non-zero with message on stderr otherwise.
"""
import sys

def main():
    if len(sys.argv) < 4:
        sys.stderr.write("usage: merge_docx.py host annex1 [annex2 ...] output\n")
        return 2
    host_path = sys.argv[1]
    out_path = sys.argv[-1]
    annex_paths = sys.argv[2:-1]
    from docxcompose.composer import Composer
    from docx import Document
    from docx.enum.text import WD_BREAK
    master = Document(host_path)
    composer = Composer(master)
    for annex_path in annex_paths:
        # Page break BEFORE each appended annex, so it starts on a fresh page.
        brk = master.add_paragraph()
        brk.add_run().add_break(WD_BREAK.PAGE)
        composer.append(Document(annex_path))
    composer.save(out_path)
    print("merged ok (%d annexes, page break before each)" % len(annex_paths))
    return 0

if __name__ == "__main__":
    sys.exit(main())
